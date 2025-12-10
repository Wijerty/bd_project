from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_code_block(document, code_text):
    """Adds a code block with a border and monospaced font."""
    table = document.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

def add_data_table(document, headers, data):
    """Adds a table with headers and data to the document."""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Make header bold
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    for row_data in data:
        row_cells = table.add_row().cells
        for i, item in enumerate(row_data):
            row_cells[i].text = str(item)
    document.add_paragraph() # Add spacing after table

def create_report():
    document = Document()

    # --- Title ---
    title = document.add_heading('Отчет о проекте "Антифрод система для P2P переводов"', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 1. Introduction ---
    document.add_heading('1. Введение', level=1)
    document.add_paragraph('Проект представляет собой специализированную систему для выявления и предотвращения мошеннических действий в сфере P2P (peer-to-peer) переводов. Система анализирует транзакции в реальном времени, используя поведенческие и сетевые признаки для оценки рисков.')
    
    document.add_heading('Цель проекта', level=2)
    document.add_paragraph('Создание надежного механизма для выявления подозрительных транзакций, защиты средств клиентов и минимизации финансовых потерь от мошенничества.')

    # --- 2. Domain Description ---
    document.add_heading('2. Описание предметной области и сущностей', level=1)
    document.add_paragraph('База данных спроектирована для хранения полной информации о финансовой деятельности и цифровом следе пользователей. Основные сущности системы:')

    entities = [
        ('Клиент (Client)', 'Пользователь системы. Хранит персональные данные (ФИО, телефон, email), статус KYC, уровень риска и статус блокировки.'),
        ('Счёт/Карта (Account/Card)', 'Финансовые инструменты клиента. Включает тип счета (карта, банковский счет, кошелек), валюту, баланс и лимиты (дневные/месячные).'),
        ('Транзакция (Transaction)', 'Запись о переводе средств. Содержит сумму, валюту, дату, статус, а также ссылки на устройство и IP-адрес инициатора.'),
        ('Устройство (Device)', 'Информация об устройстве, с которого совершается операция (fingerprint, тип, ОС, браузер, user-agent).'),
        ('IP-адрес (IP Address)', 'Сетевой адрес инициатора транзакции, включая геолокацию (страна, город, координаты), провайдера и признаки использования прокси/Tor.'),
        ('Правила/Алёрты (Rules/Alerts)', 'Набор критериев для оценки подозрительности транзакций с весовыми коэффициентами.'),
        ('Чёрные списки (Blacklists)', 'Списки заблокированных сущностей (клиентов, карт, IP, устройств) с указанием причины и срока блокировки.'),
        ('Связи (Client Relationships)', 'Граф социальных и финансовых связей между клиентами («кто кому переводит»), позволяющий выявлять скрытые группы.')
    ]

    for name, desc in entities:
        p = document.add_paragraph()
        run = p.add_run(name + ': ')
        run.bold = True
        p.add_run(desc)

    # --- 3. Data Dictionary ---
    document.add_heading('3. Полный словарь данных (Data Dictionary)', level=1)
    document.add_paragraph('Ниже приведен полный список всех атрибутов (полей) для каждой сущности базы данных.')

    tables_structure = [
        ('Client (Клиенты)', [
            ('client_id', 'SERIAL PK', 'Уникальный идентификатор'),
            ('first_name', 'VARCHAR(100)', 'Имя'),
            ('last_name', 'VARCHAR(100)', 'Фамилия'),
            ('date_of_birth', 'DATE', 'Дата рождения'),
            ('phone_number', 'VARCHAR(20)', 'Телефон'),
            ('email', 'VARCHAR(255)', 'Email'),
            ('registration_date', 'TIMESTAMP', 'Дата регистрации'),
            ('kyc_status', 'VARCHAR(50)', 'Статус верификации'),
            ('risk_level', 'DECIMAL(3,2)', 'Уровень риска (0.0-1.0)'),
            ('is_blocked', 'BOOLEAN', 'Заблокирован ли клиент')
        ]),
        ('Account (Счета)', [
            ('account_id', 'SERIAL PK', 'Уникальный идентификатор'),
            ('client_id', 'INTEGER FK', 'Ссылка на клиента'),
            ('account_number', 'VARCHAR(50)', 'Номер счета/карты'),
            ('account_type', 'VARCHAR(20)', 'Тип (card, bank_account)'),
            ('currency', 'CHAR(3)', 'Валюта (RUB, USD)'),
            ('balance', 'DECIMAL(15,2)', 'Текущий баланс'),
            ('opening_date', 'TIMESTAMP', 'Дата открытия'),
            ('is_active', 'BOOLEAN', 'Активен ли счет'),
            ('daily_limit', 'DECIMAL', 'Дневной лимит'),
            ('monthly_limit', 'DECIMAL', 'Месячный лимит')
        ]),
        ('Transaction (Транзакции)', [
            ('transaction_id', 'SERIAL PK', 'Уникальный идентификатор'),
            ('sender_account_id', 'INTEGER FK', 'Счет отправителя'),
            ('receiver_account_id', 'INTEGER FK', 'Счет получателя'),
            ('amount', 'DECIMAL(15,2)', 'Сумма'),
            ('currency', 'CHAR(3)', 'Валюта'),
            ('transaction_date', 'TIMESTAMP', 'Время транзакции'),
            ('transaction_type', 'VARCHAR(20)', 'Тип (P2P, merchant)'),
            ('status', 'VARCHAR(20)', 'Статус (completed, failed)'),
            ('ip_address_id', 'INTEGER FK', 'IP адрес'),
            ('device_id', 'INTEGER FK', 'Устройство'),
            ('location_coordinates', 'POINT', 'Гео-координаты'),
            ('fraud_score', 'DECIMAL', 'Оценка фрода'),
            ('is_flagged', 'BOOLEAN', 'Помечена как подозрительная'),
            ('flagged_reason', 'TEXT', 'Причина метки')
        ]),
        ('Device (Устройства)', [
            ('device_id', 'SERIAL PK', 'ID устройства'),
            ('device_fingerprint', 'VARCHAR', 'Уникальный отпечаток'),
            ('device_type', 'VARCHAR', 'Тип (mobile, desktop)'),
            ('os', 'VARCHAR', 'Операционная система'),
            ('browser', 'VARCHAR', 'Браузер'),
            ('user_agent', 'TEXT', 'User-Agent строка'),
            ('first_seen_date', 'TIMESTAMP', 'Впервые замечено'),
            ('risk_score', 'DECIMAL', 'Уровень риска устройства')
        ]),
        ('IPAddress (IP-адреса)', [
            ('ip_address_id', 'SERIAL PK', 'ID адреса'),
            ('ip_address', 'INET', 'IP адрес'),
            ('country', 'VARCHAR', 'Страна'),
            ('city', 'VARCHAR', 'Город'),
            ('isp', 'VARCHAR', 'Провайдер'),
            ('is_proxy', 'BOOLEAN', 'Является ли прокси'),
            ('is_tor', 'BOOLEAN', 'Является ли TOR узлом'),
            ('risk_score', 'DECIMAL', 'Уровень риска IP')
        ])
    ]

    for table_name, fields in tables_structure:
        document.add_heading(table_name, level=2)
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Поле'
        hdr_cells[1].text = 'Тип'
        hdr_cells[2].text = 'Описание'
        
        for field in fields:
            row_cells = table.add_row().cells
            row_cells[0].text = field[0]
            row_cells[1].text = field[1]
            row_cells[2].text = field[2]
        document.add_paragraph()

    # --- 4. Database Visualization ---
    document.add_heading('4. Визуализация структуры базы данных', level=1)
    document.add_paragraph('Ниже представлена схема связей между основными таблицами базы данных.')

    # Check for images and insert them if they exist
    high_level_img = 'high_level_diagram.png'
    er_diagram_img = 'er_diagram.png'

    if os.path.exists(high_level_img):
        document.add_heading('4.1. Логическая структура (High-Level)', level=2)
        document.add_picture(high_level_img, width=Inches(6))
        document.add_paragraph('Рис. 1. Высокоуровневая архитектура данных.')
    
    if os.path.exists(er_diagram_img):
        document.add_heading('4.2. Детальная ER-диаграмма', level=2)
        document.add_picture(er_diagram_img, width=Inches(6))
        document.add_paragraph('Рис. 2. Детальная схема базы данных.')

    if not os.path.exists(high_level_img) and not os.path.exists(er_diagram_img):
        document.add_paragraph('[Место для вставки скриншотов из файла DB_VISUALIZATION.html]')
        
        # Fallback to text diagram
        document.add_paragraph('Текстовое представление схемы:')
        erd_diagram = """
       +----------------+          +-----------------+
       |     Client     | 1 ---- N |     Account     |
       +----------------+          +-----------------+
               |                            | 1
               |                            |
               | N                          | N
       +----------------+          +-----------------+          +----------------+
       | ClientRelation |          |   Transaction   | N ---- 1 |     Device     |
       +----------------+          +-----------------+          +----------------+
                                            |
                                            | N
                                            |
                                   +-----------------+
                                   |    IPAddress    |
                                   +-----------------+
        """
        add_code_block(document, erd_diagram)

    document.add_paragraph('Ключевые связи:')
    relationships = [
        'Один Клиент может иметь несколько Счетов (1:N).',
        'Один Счет может участвовать во множестве Транзакций как отправитель или получатель (1:N).',
        'Каждая Транзакция связана с одним Устройством и одним IP-адресом (N:1).',
        'Клиенты могут иметь связи друг с другом через таблицу ClientRelationships (N:N логически).',
    ]
    for rel in relationships:
        document.add_paragraph(rel, style='List Bullet')

    # --- 5. Architecture ---
    document.add_heading('5. Архитектура системы', level=1)
    document.add_paragraph('Проект реализован как микросервисная архитектура, разворачиваемая с помощью Docker.')

    components = [
        ('База данных (PostgreSQL)', 'Центральное хранилище данных. Используется для хранения транзакций, профилей клиентов и выполнения сложных аналитических запросов (включая рекурсивные CTE).'),
        ('Backend аналитики (Python)', 'Скрипты (fraud_detection.py, advanced_fraud_detection.py) реализуют логику скоринга, выявления паттернов ("карусели", всплески) и взаимодействия с БД.'),
        ('Веб-интерфейс (Security Dashboard)', 'Приложение на Flask, предоставляющее аналитикам визуализацию метрик безопасности, тепловые карты активности и управление инцидентами.'),
        ('Генератор данных', 'Скрипт transaction_generator.py для создания синтетических данных, позволяющий тестировать систему под нагрузкой.')
    ]

    for name, desc in components:
        p = document.add_paragraph()
        run = p.add_run(name + ': ')
        run.bold = True
        p.add_run(desc)

    # --- 6. Implementation Details (Code Snippets) ---
    document.add_heading('6. Детали реализации и примеры кода', level=1)
    
    document.add_heading('6.1. Структура таблицы транзакций (SQL)', level=2)
    document.add_paragraph('Таблица Transaction является центральным элементом системы. Ниже приведен SQL-код её создания:')
    
    sql_code = """CREATE TABLE Transaction (
    transaction_id SERIAL PRIMARY KEY,
    sender_account_id INTEGER NOT NULL,
    receiver_account_id INTEGER NOT NULL,
    amount DECIMAL(15,2) NOT NULL,
    currency CHAR(3) DEFAULT 'RUB',
    transaction_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    transaction_type VARCHAR(20),
    status VARCHAR(20),
    ip_address_id INTEGER,
    device_id INTEGER,
    location_coordinates POINT,
    fraud_score DECIMAL(5,2) DEFAULT 0.0,
    is_flagged BOOLEAN DEFAULT FALSE,
    flagged_reason TEXT,
    FOREIGN KEY (sender_account_id) REFERENCES Account(account_id),
    FOREIGN KEY (receiver_account_id) REFERENCES Account(account_id),
    FOREIGN KEY (ip_address_id) REFERENCES IPAddress(ip_address_id),
    FOREIGN KEY (device_id) REFERENCES Device(device_id)
);"""
    add_code_block(document, sql_code)

    document.add_heading('6.2. Обнаружение "Каруселей" (Python + SQL)', level=2)
    document.add_paragraph('Для выявления сложных схем отмывания денег используется рекурсивный CTE запрос, который строит цепочки транзакций. Реализация в advanced_fraud_detection.py:')

    python_carousel_code = """def detect_carousel_patterns(self, time_window_hours=24) -> List[Dict]:
    query = \"\"\"
    WITH RECURSIVE transaction_paths AS (
        SELECT 
            t.sender_account_id,
            t.receiver_account_id,
            t.amount,
            ARRAY[t.sender_account_id, t.receiver_account_id] as account_path,
            1 as path_length
        FROM transaction t
        WHERE t.transaction_date >= NOW() - INTERVAL '%s hours'
        
        UNION ALL
        
        SELECT 
            tp.sender_account_id,
            t.receiver_account_id,
            tp.amount + t.amount,
            tp.account_path || t.receiver_account_id,
            tp.path_length + 1
        FROM transaction_paths tp
        JOIN transaction t ON tp.receiver_account_id = t.sender_account_id
        WHERE t.transaction_date >= NOW() - INTERVAL '%s hours'
        AND tp.path_length < 6
        AND t.receiver_account_id != ALL(tp.account_path[1:-1])
    )
    SELECT * FROM transaction_paths
    WHERE account_path[1] = account_path[array_length(account_path, 1)]
    AND path_length >= 3;
    \"\"\"
    # ... execution logic ..."""
    add_code_block(document, python_carousel_code)

    document.add_heading('6.3. Проверка новых устройств', level=2)
    document.add_paragraph('Система отслеживает использование новых устройств. Если устройство замечено впервые за последние 24 часа, это повышает риск транзакции.')
    
    device_check_code = """def detect_new_devices(self, hours: int = 24):
    query = \"\"\"
        SELECT t.transaction_id, t.transaction_date, d.device_fingerprint
        FROM Transaction t
        JOIN Device d ON t.device_id = d.device_id
        WHERE d.first_seen_date >= NOW() - INTERVAL '%s hours'
    \"\"\"
    # ... execution logic ..."""
    add_code_block(document, device_check_code)

    # --- 7. Fraud Detection Algorithms ---
    document.add_heading('7. Алгоритмы выявления мошенничества', level=1)
    document.add_paragraph('Система реализует ряд ключевых аналитических запросов для детектирования фрода:')

    algorithms = [
        ('Всплески частоты переводов', 'Анализ аномальной активности по счету за короткий промежуток времени. Реализуется через агрегацию транзакций по временным окнам и сравнение с историческим профилем.'),
        ('«Карусели» переводов', 'Выявление циклических схем отмывания денег (A -> B -> C -> A). Используются рекурсивные SQL-запросы (CTE) для построения цепочек транзакций глубиной до 10 уровней.'),
        ('Транзакции с новыми устройствами', 'Отслеживание операций с устройств, ранее не ассоциированных с клиентом. Если устройство замечено впервые менее 24 часов назад, транзакция помечается как рискованная.'),
        ('Графовые запросы на кластеры связности', 'Анализ социальных графов для выявления организованных преступных групп ("дроп-ферм"). Поиск компонент связности в графе отношений клиентов.'),
        ('Скоринг по правилам и весам', 'Комплексная оценка риска. Каждое сработавшее правило (например, "Сумма > 10000", "Рискованный IP") добавляет вес к общему fraud_score. Превышение порога ведет к блокировке.')
    ]

    for name, desc in algorithms:
        document.add_heading(name, level=2)
        document.add_paragraph(desc)

    # --- 8. Visualization (Dashboard) ---
    document.add_heading('8. Визуализация (Security Dashboard)', level=1)
    document.add_paragraph('Для операторов безопасности разработан веб-интерфейс, отображающий:')
    
    dashboard_features = [
        'Статистику по заблокированным транзакциям и общим объемам.',
        'Тепловые карты активности по регионам и времени.',
        'Детальную информацию по конкретным инцидентам и профилям клиентов.',
        'Интерфейс управления черными списками и правилами скоринга.'
    ]
    
    for feature in dashboard_features:
        document.add_paragraph(feature, style='List Bullet')

    # --- 9. Sample Data ---
    document.add_heading('9. Примеры данных в таблицах', level=1)
    document.add_paragraph('Ниже приведены примеры данных, хранящихся в основных таблицах системы.')

    # Client Data
    document.add_heading('9.1. Таблица Client (Клиенты)', level=2)
    client_headers = ['ID', 'Имя', 'Фамилия', 'Телефон', 'Email', 'Статус KYC', 'Риск']
    client_data = [
        ['1', 'Ivan', 'Petrov', '+79991234567', 'ivan.petrov@email.com', 'verified', '0.1'],
        ['2', 'Maria', 'Sidorova', '+79997654321', 'maria.sidorova@email.com', 'verified', '0.05'],
        ['3', 'Alexey', 'Ivanov', '+79991112233', 'alexey.ivanov@email.com', 'pending', '0.3']
    ]
    add_data_table(document, client_headers, client_data)

    # Account Data
    document.add_heading('9.2. Таблица Account (Счета)', level=2)
    account_headers = ['ID', 'Client ID', 'Номер счета', 'Тип', 'Баланс']
    account_data = [
        ['1', '1', '4276123456789012', 'card', '50000.00'],
        ['2', '2', '4276987654321098', 'card', '30000.00'],
        ['3', '3', '4276112233445566', 'card', '15000.00']
    ]
    add_data_table(document, account_headers, account_data)

    # Transaction Data
    document.add_heading('9.3. Таблица Transaction (Транзакции)', level=2)
    trans_headers = ['ID', 'От кого', 'Кому', 'Сумма', 'Тип', 'Статус', 'Device ID']
    trans_data = [
        ['1', '1', '2', '5000.00', 'P2P', 'completed', '1'],
        ['2', '2', '3', '3000.00', 'P2P', 'completed', '2'],
        ['3', '3', '1', '1000.00', 'P2P', 'completed', '1']
    ]
    add_data_table(document, trans_headers, trans_data)

    # Device Data
    document.add_heading('9.4. Таблица Device (Устройства)', level=2)
    device_headers = ['ID', 'Fingerprint', 'Тип', 'OS', 'Browser']
    device_data = [
        ['1', 'fingerprint_001', 'mobile', 'Android', 'Chrome'],
        ['2', 'fingerprint_002', 'desktop', 'Windows', 'Firefox']
    ]
    add_data_table(document, device_headers, device_data)

    # IP Data
    document.add_heading('9.5. Таблица IPAddress (IP-адреса)', level=2)
    ip_headers = ['ID', 'IP', 'Страна', 'Город', 'Провайдер']
    ip_data = [
        ['1', '192.168.1.100', 'Russia', 'Moscow', 'MTS'],
        ['2', '192.168.1.101', 'Russia', 'Saint Petersburg', 'Beeline']
    ]
    add_data_table(document, ip_headers, ip_data)

    # Rules Data
    document.add_heading('9.6. Таблица Rule (Правила)', level=2)
    rule_headers = ['Название', 'Описание', 'Условие', 'Вес']
    rule_data = [
        ['HighAmount', 'Сумма выше порога', 'amount > 10000', '5.0'],
        ['NewDevice', 'Новое устройство', 'device_first_seen < 24h', '3.0'],
        ['HighRiskIP', 'Рискованный IP', 'ip_risk_score > 0.8', '4.0']
    ]
    add_data_table(document, rule_headers, rule_data)

    # --- 10. Conclusion ---
    document.add_heading('10. Заключение', level=1)
    document.add_paragraph('Разработанная система покрывает основные векторы атак в P2P переводах. Использование реляционной СУБД PostgreSQL с поддержкой рекурсивных запросов позволяет эффективно выявлять сложные схемы мошенничества (графы, цепочки) без необходимости внедрения специализированных графовых баз данных на начальном этапе. Модульная архитектура обеспечивает масштабируемость и удобство поддержки.')

    # Save document
    filename = 'Отчет_Антифрод_Система_Final_v3.docx'
    try:
        document.save(filename)
        print(f"Report generated successfully: {filename}")
    except PermissionError:
        print(f"Error: Could not save '{filename}'. Please close the file if it is open in Word and try again.")

if __name__ == "__main__":
    create_report()
